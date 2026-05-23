"""
Report generator.

Produces two outputs from a ReviewResult:
  1. A Word document (.docx) — for download
  2. A Markdown string — for in-app display

Both follow the same structure (as specified in Section 9 of the HLD design doc):

  - Summary (counts, top concerns, one-line status per standard)
  - Findings by standard (every finding with full evidence)
  - Severity legend

Brand: Accenture purple for headings, table headers, severity callouts.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

from .models import ReviewResult, StandardResult, Finding


# ---------------------------------------------------------------------------
# Colour palette (Accenture)
# ---------------------------------------------------------------------------
COL_PURPLE = RGBColor(0xA1, 0x00, 0xFF)
COL_PURPLE_DEEP = RGBColor(0x7A, 0x00, 0xC2)
COL_BLACK = RGBColor(0x00, 0x00, 0x00)
COL_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COL_MUTED = RGBColor(0x52, 0x52, 0x52)
COL_HIGH = RGBColor(0xBE, 0x12, 0x3C)
COL_MEDIUM = RGBColor(0xB4, 0x53, 0x09)
COL_LOW = RGBColor(0x0F, 0x76, 0x6E)

# Hex strings for table shading
HEX_PURPLE_DEEP = "7A00C2"
HEX_PURPLE_WASH = "F4E5FF"
HEX_TINT = "FAFAFA"


SEVERITY_COLORS = {
    "High": COL_HIGH,
    "Medium": COL_MEDIUM,
    "Low": COL_LOW,
}


# ===========================================================================
# Public entry points
# ===========================================================================

def build_markdown_report(review: ReviewResult) -> str:
    """Render the review as Markdown for in-app display."""
    lines: List[str] = []

    lines.append(f"# Pre-Review Report")
    lines.append(f"**Submission:** {review.hld_filename}  ")
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  ")
    lines.append(f"**Agent model:** `{review.model}`  ")
    lines.append("")

    # Summary
    lines.append("## Summary")
    lines.append("")
    high = review.findings_by_severity("High")
    medium = review.findings_by_severity("Medium")
    low = review.findings_by_severity("Low")
    lines.append(f"- **Total findings:** {review.total_findings}")
    lines.append(f"- **High severity:** {high}")
    lines.append(f"- **Medium severity:** {medium}")
    lines.append(f"- **Low severity:** {low}")
    lines.append("")

    # Status per standard
    lines.append("### Status per standard")
    lines.append("")
    lines.append("| # | Standard | Findings | Status |")
    lines.append("|---|---|---|---|")
    for r in review.results:
        status = _standard_status(r)
        lines.append(
            f"| {r.standard.display_id} | {r.standard.name} | {len(r.findings)} | {status} |"
        )
    lines.append("")

    # Findings by standard
    lines.append("## Findings by standard")
    lines.append("")

    for r in review.results:
        lines.append(f"### {r.standard.display_id} — {r.standard.name}")
        lines.append("")
        if r.error:
            lines.append(f"> ⚠️ Standard could not be evaluated: {r.error}")
            lines.append("")
            continue

        if not r.findings:
            lines.append("_No findings. Standard is satisfied or not applicable to this design._")
            lines.append("")
            continue

        for idx, f in enumerate(r.findings, start=1):
            finding_id = f"F-{r.standard.display_id}.{idx}"
            lines.append(f"#### {finding_id} · **{f.severity}** · {f.check_id} {f.check_name}")
            lines.append("")
            lines.append(f"**Issue.** {f.issue}")
            lines.append("")
            lines.append(f"**Evidence.** {f.evidence_section}")
            lines.append("")
            lines.append(f"> {_quote_for_md(f.evidence_quote)}")
            lines.append("")
            lines.append(f"**Authority.** {f.authority}")
            lines.append("")
            lines.append(f"**Recommendation.** {f.recommendation}")
            lines.append("")

    # Legend
    lines.append("---")
    lines.append("")
    lines.append("**Severity legend.** ")
    lines.append("**High** — address before AAG review. Otherwise likely conditional approval or return for revision.  ")
    lines.append("**Medium** — should be addressed or explicitly accepted with rationale.  ")
    lines.append("**Low** — clarification opportunity; unlikely to block approval.")

    return "\n".join(lines)


def build_docx_report(review: ReviewResult) -> bytes:
    """Render the review as a Word document and return the bytes."""
    doc = Document()
    _setup_document(doc)

    _add_cover(doc, review)
    _add_summary(doc, review)
    _add_findings(doc, review)
    _add_legend(doc)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# Doc setup
# ===========================================================================

def _setup_document(doc: Document):
    """Set Arial as the default font, US Letter page size, 1" margins."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = COL_TEXT


# ===========================================================================
# Sections
# ===========================================================================

def _add_cover(doc: Document, review: ReviewResult):
    # Eyebrow
    p = doc.add_paragraph()
    run = p.add_run("ARCHITECTURE GOVERNANCE  ·  PRE-REVIEW REPORT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COL_PURPLE

    # Title
    p = doc.add_paragraph()
    run = p.add_run("Pre-Review Report")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(32)
    run.font.color.rgb = COL_BLACK

    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run(f"Submission: {review.hld_filename}")
    run.font.size = Pt(14)
    run.font.color.rgb = COL_MUTED

    # Horizontal accent rule
    _add_horizontal_rule(doc, COL_PURPLE)

    # Metadata box
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False
    _set_table_col_widths(table, [Inches(2.0), Inches(4.5)])
    _style_metadata_table(table, [
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Agent model", review.model),
        ("Total findings", str(review.total_findings)),
    ])

    doc.add_paragraph()


def _add_summary(doc: Document, review: ReviewResult):
    _add_heading(doc, "Summary", level=1)

    # Severity counts in a grid
    high = review.findings_by_severity("High")
    medium = review.findings_by_severity("Medium")
    low = review.findings_by_severity("Low")

    p = doc.add_paragraph()
    run = p.add_run("Total findings: ")
    run.bold = True
    p.add_run(str(review.total_findings))

    # 3-cell severity table
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    col_w = Inches(2.16)
    _set_table_col_widths(table, [col_w, col_w, col_w])

    headers = [
        ("HIGH", COL_HIGH),
        ("MEDIUM", COL_MEDIUM),
        ("LOW", COL_LOW),
    ]
    counts = [high, medium, low]
    for ci, (label, colour) in enumerate(headers):
        cell = table.rows[0].cells[ci]
        _set_cell_shading(cell, HEX_TINT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colour
        _set_cell_borders(cell, color="D9D9D9")

    for ci, count in enumerate(counts):
        cell = table.rows[1].cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(count))
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = COL_BLACK
        _set_cell_borders(cell, color="D9D9D9")

    doc.add_paragraph()

    # Per-standard status table
    _add_heading(doc, "Status per standard", level=2)
    _add_status_table(doc, review)
    doc.add_paragraph()


def _add_status_table(doc: Document, review: ReviewResult):
    table = doc.add_table(rows=1 + len(review.results), cols=4)
    table.autofit = False
    widths = [Inches(0.5), Inches(3.5), Inches(0.9), Inches(1.6)]
    _set_table_col_widths(table, widths)

    # Header row
    headers = ["#", "Standard", "Findings", "Status"]
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        _set_cell_shading(cell, HEX_PURPLE_DEEP)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_borders(cell, color="D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, r in enumerate(review.results, start=1):
        row = table.rows[ri]
        for ci, val in enumerate([
            r.standard.display_id,
            r.standard.name,
            str(len(r.findings)),
            _standard_status(r),
        ]):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = COL_TEXT
            _set_cell_borders(cell, color="D9D9D9")


def _add_findings(doc: Document, review: ReviewResult):
    doc.add_page_break()
    _add_heading(doc, "Findings by standard", level=1)

    for r in review.results:
        _add_heading(doc, f"{r.standard.display_id} — {r.standard.name}", level=2)

        # Standard purpose
        p = doc.add_paragraph()
        run = p.add_run(r.standard.purpose)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = COL_MUTED

        if r.error:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ Standard could not be evaluated: {r.error}")
            run.font.color.rgb = COL_HIGH
            continue

        if not r.findings:
            p = doc.add_paragraph()
            run = p.add_run(
                "No findings. Standard is satisfied or not applicable to this design."
            )
            run.italic = True
            run.font.color.rgb = COL_MUTED
            continue

        for idx, f in enumerate(r.findings, start=1):
            finding_id = f"F-{r.standard.display_id}.{idx}"
            _add_finding(doc, finding_id, f)


def _add_finding(doc: Document, finding_id: str, f: Finding):
    """Render a single finding as a small 2-column table."""
    # Header line: ID · severity · check
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(finding_id)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COL_BLACK

    run = p.add_run("  ·  ")
    run.font.color.rgb = COL_MUTED

    sev_colour = SEVERITY_COLORS.get(f.severity, COL_TEXT)
    run = p.add_run(f.severity.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = sev_colour

    run = p.add_run("  ·  ")
    run.font.color.rgb = COL_MUTED

    run = p.add_run(f"{f.check_id} {f.check_name}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COL_BLACK

    # 2-column table with field labels and values
    rows = [
        ("Issue", f.issue),
        ("Evidence", f.evidence_section),
        ("Quote", f.evidence_quote),
        ("Authority", f.authority),
        ("Recommendation", f.recommendation),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    _set_table_col_widths(table, [Inches(1.3), Inches(5.2)])

    for ri, (label, value) in enumerate(rows):
        # Label cell
        cell = table.rows[ri].cells[0]
        _set_cell_shading(cell, HEX_PURPLE_WASH)
        _set_cell_borders(cell, color="ECECEC")
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COL_PURPLE_DEEP

        # Value cell
        cell = table.rows[ri].cells[1]
        _set_cell_borders(cell, color="ECECEC")
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = COL_TEXT
        # Italics for the quote
        if label == "Quote":
            run.italic = True
            run.font.color.rgb = COL_MUTED


def _add_legend(doc: Document):
    doc.add_page_break()
    _add_heading(doc, "Severity legend", level=1)

    legend = [
        ("High", COL_HIGH,
         "Address before AAG review. Failure to address would normally result in conditional approval or return for revision."),
        ("Medium", COL_MEDIUM,
         "Should be addressed or explicitly accepted with rationale. The AAG will likely raise these in the meeting if unresolved."),
        ("Low", COL_LOW,
         "Clarification opportunity. Worth resolving but unlikely to block approval."),
    ]

    table = doc.add_table(rows=len(legend), cols=2)
    table.autofit = False
    _set_table_col_widths(table, [Inches(1.2), Inches(5.3)])

    for ri, (label, colour, desc) in enumerate(legend):
        cell = table.rows[ri].cells[0]
        _set_cell_borders(cell, color="ECECEC")
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colour

        cell = table.rows[ri].cells[1]
        _set_cell_borders(cell, color="ECECEC")
        p = cell.paragraphs[0]
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = COL_TEXT

    # Disclosure
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "This report is generated by the AI Pre-Review Agent. It surfaces concerns "
        "against universal architectural standards for the AAG to evaluate. The "
        "agent does not approve or reject designs. All findings are advisory."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COL_MUTED


# ===========================================================================
# Low-level helpers
# ===========================================================================

def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COL_BLACK
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COL_PURPLE_DEEP
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = COL_BLACK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)


def _add_horizontal_rule(doc: Document, color: RGBColor):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A100FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "D9D9D9", size: str = "4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_table_col_widths(table, widths):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]


def _style_metadata_table(table, rows):
    for ri, (label, value) in enumerate(rows):
        cell = table.rows[ri].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COL_MUTED
        _set_cell_borders(cell, color="ECECEC")

        cell = table.rows[ri].cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = COL_TEXT
        _set_cell_borders(cell, color="ECECEC")


def _standard_status(r: StandardResult) -> str:
    if r.error:
        return "Error"
    high = sum(1 for f in r.findings if f.severity == "High")
    if high:
        return f"⚠ {high} high"
    if r.findings:
        return f"{len(r.findings)} findings"
    return "Clear"


def _quote_for_md(text: str) -> str:
    """Format a quote for Markdown blockquote. Collapse newlines."""
    return text.replace("\n", " ").strip()[:400]
