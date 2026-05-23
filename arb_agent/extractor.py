"""
HLD extractor.

Reads a High-Level Design document (.docx or .pdf) and produces a structured
text representation the agent can reason over. The goal is not a perfect parse —
the goal is to preserve enough section structure that the agent can cite
"Section X.Y" when it finds something.

For .docx we use python-docx and walk paragraphs and tables in document order.
For .pdf we use pdfplumber and extract text page by page.
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import List, Optional


@dataclass
class HLDDocument:
    filename: str
    text: str                # full document text, with section headings inlined
    section_count: int       # rough count of headings found
    page_count: Optional[int] = None
    table_count: int = 0
    word_count: int = 0


def extract_from_bytes(filename: str, data: bytes) -> HLDDocument:
    """Dispatch on file extension and extract the HLD content."""
    name = filename.lower()
    if name.endswith(".docx"):
        return _extract_docx(filename, data)
    elif name.endswith(".pdf"):
        return _extract_pdf(filename, data)
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Supported: .docx, .pdf"
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(filename: str, data: bytes) -> HLDDocument:
    """Extract structured text from a .docx file.

    Walks the document body. For each paragraph: if it's a heading, prefix it
    with markdown-style #s so the agent can see hierarchy. For each table:
    flatten into a simple text representation.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(BytesIO(data))

    out_lines: List[str] = []
    section_count = 0
    table_count = 0

    # Iterate over body in document order — paragraphs and tables interleave.
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            # It's a paragraph
            para = next((p for p in doc.paragraphs if p._element is child), None)
            if para is None:
                continue
            text = (para.text or "").strip()
            if not text:
                continue

            # Detect heading by style name (e.g. "Heading 1", "Heading 2")
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading"):
                # Extract heading level (default to 2)
                level = 2
                for tok in style_name.split():
                    if tok.isdigit():
                        level = int(tok)
                        break
                prefix = "#" * max(1, min(level, 6))
                out_lines.append(f"\n{prefix} {text}\n")
                section_count += 1
            elif style_name in ("title",):
                out_lines.append(f"\n# {text}\n")
                section_count += 1
            else:
                out_lines.append(text)

        elif tag == qn("w:tbl"):
            # It's a table — find matching Table object
            tbl = next((t for t in doc.tables if t._element is child), None)
            if tbl is None:
                continue
            table_count += 1
            out_lines.append(_render_docx_table(tbl))

    full_text = "\n".join(out_lines).strip()
    word_count = len(full_text.split())

    return HLDDocument(
        filename=filename,
        text=full_text,
        section_count=section_count,
        table_count=table_count,
        word_count=word_count,
    )


def _render_docx_table(tbl) -> str:
    """Flatten a docx table into a readable text block."""
    rows_out: List[str] = []
    for row in tbl.rows:
        cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
        rows_out.append(" | ".join(cells))
    if not rows_out:
        return ""
    return "\n[TABLE]\n" + "\n".join(rows_out) + "\n[/TABLE]\n"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _extract_pdf(filename: str, data: bytes) -> HLDDocument:
    """Extract text from a .pdf using pdfplumber.

    PDFs don't carry heading metadata reliably. We extract page-by-page text
    and let the agent's prompt do the section recognition. We also extract
    tables when pdfplumber detects them.
    """
    import pdfplumber

    out_lines: List[str] = []
    table_count = 0
    page_count = 0

    with pdfplumber.open(BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                out_lines.append(f"\n[Page {idx}]\n{text}")

            # Tables
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for tbl in tables:
                table_count += 1
                rendered = _render_pdf_table(tbl)
                if rendered:
                    out_lines.append(rendered)

    full_text = "\n".join(out_lines).strip()
    # Rough heading detection for stats — count lines that look like headings
    section_count = sum(
        1
        for line in full_text.splitlines()
        if line.strip() and _looks_like_heading(line.strip())
    )
    word_count = len(full_text.split())

    return HLDDocument(
        filename=filename,
        text=full_text,
        section_count=section_count,
        page_count=page_count,
        table_count=table_count,
        word_count=word_count,
    )


def _render_pdf_table(tbl) -> str:
    rows_out: List[str] = []
    for row in tbl:
        cells = [(c or "").strip().replace("\n", " ") for c in row]
        rows_out.append(" | ".join(cells))
    if not rows_out:
        return ""
    return "\n[TABLE]\n" + "\n".join(rows_out) + "\n[/TABLE]\n"


def _looks_like_heading(line: str) -> bool:
    """Cheap heuristic for heading-like lines in PDF text."""
    if len(line) > 120:
        return False
    if line.endswith("."):
        return False
    # Starts with section number like "1.", "1.1", "1.1.1"
    parts = line.split()
    if not parts:
        return False
    first = parts[0]
    if first[0].isdigit() and "." in first:
        return True
    # Or short Title Case line
    words = parts
    if 2 <= len(words) <= 12 and sum(1 for w in words if w[:1].isupper()) >= len(words) // 2:
        return True
    return False
